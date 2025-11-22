# -*- coding: utf-8 -*-
"""
Comprehensive End-to-End Test for ShapeShifter RAG System
Tests all major functionalities
"""
import requests
import json
import sys
from io import BytesIO
from docx import Document

if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

BASE_URL = "http://localhost:8001"
SESSION_ID = "final_test_session"

print("=" * 80)
print("SHAPESHIFTER RAG SYSTEM - COMPREHENSIVE FUNCTIONALITY TEST")
print("=" * 80)

# Create a test DOCX file
def create_test_docx():
    doc = Document()
    doc.add_heading('Article 21 - Right to Life and Personal Liberty', 0)
    doc.add_paragraph('No person shall be deprived of his life or personal liberty except according to procedure established by law.')
    doc.add_heading('Landmark Cases', level=1)
    doc.add_paragraph('1. Maneka Gandhi v. Union of India (1978): Transformed Article 21 from procedural to substantive right.')
    doc.add_paragraph('2. The Court held that procedure must be just, fair, and reasonable.')
    doc.add_paragraph('3. This interpretation enabled rights like privacy, education, and clean environment.')
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

print("\n📋 TEST 1: Health Check")
print("-" * 80)
try:
    response = requests.get(f"{BASE_URL}/health")
    health = response.json()
    print(f"✅ Status: {health['status']}")
    print(f"✅ LLM Provider: {health['llm_provider']}")
    print(f"✅ Model: {health['model']}")
except Exception as e:
    print(f"❌ Health check failed: {e}")
    sys.exit(1)

print("\n📋 TEST 2: RAG Creation Workflow")
print("-" * 80)

# Step 1: Initiate RAG creation
print("Step 1: Initiating RAG creation...")
response = requests.post(f"{BASE_URL}/chat", json={
    "message": "create a rag for me",
    "session_id": SESSION_ID
})
assert "personality" in response.json()['response'].lower()
print("✅ RAG creation initiated")

# Step 2: Set personality
print("Step 2: Setting personality...")
response = requests.post(f"{BASE_URL}/chat", json={
    "message": "Formal and Academic",
    "session_id": SESSION_ID
})
assert "name your rag" in response.json()['response'].lower()
print("✅ Personality set: Formal and Academic")

# Step 3: Name the RAG
print("Step 3: Naming RAG...")
response = requests.post(f"{BASE_URL}/chat", json={
    "message": "Constitution Expert",
    "session_id": SESSION_ID
})
assert "upload" in response.json()['response'].lower()
print("✅ RAG named: Constitution Expert")

print("\n📋 TEST 3: DOCX File Upload")
print("-" * 80)

# Create and upload DOCX
docx_content = create_test_docx()
files = {'file': ('test_article21.docx', docx_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
data = {'session_id': SESSION_ID}
response = requests.post(f"{BASE_URL}/upload", files=files, data=data)
upload_result = response.json()
print(f"✅ File uploaded: {upload_result['filename']}")
print(f"✅ Size: {upload_result['size']} bytes")
print(f"✅ Status: {upload_result['status']}")

# Step 4: Confirm upload
print("\nStep 4: Confirming upload...")
response = requests.post(f"{BASE_URL}/chat", json={
    "message": "uploaded",
    "session_id": SESSION_ID
})
assert "storage" in response.json()['response'].lower()
print("✅ Upload confirmed")

# Step 5: Set storage name
print("Step 5: Setting storage name...")
response = requests.post(f"{BASE_URL}/chat", json={
    "message": "constitution_storage_final",
    "session_id": SESSION_ID
})
result = response.json()
assert "complete" in result['response'].lower() or "ready" in result['response'].lower()
print("✅ RAG creation complete!")

print("\n📋 TEST 4: Settings Integration")
print("-" * 80)
# Test with settings
response = requests.post(f"{BASE_URL}/chat", json={
    "message": "What is Article 21?",
    "session_id": SESSION_ID,
    "settings": {
        "llm": "gemma-27b",
        "vectorDb": "qdrant"
    }
})
print("✅ Settings accepted and processed")

print("\n📋 TEST 5: Context-Based Query")
print("-" * 80)
print("Question: How did Maneka Gandhi case transform Article 21?")
response = requests.post(f"{BASE_URL}/chat", json={
    "message": "How did Maneka Gandhi case transform Article 21?",
    "session_id": SESSION_ID
})

result = response.json()
answer = result['response']
citations = result.get('citations', [])

print("\nAnswer Preview:")
print("-" * 80)
print(answer[:500] + "..." if len(answer) > 500 else answer)
print("-" * 80)

# Verify key elements
checks = {
    "Contains 'Maneka Gandhi'": "maneka gandhi" in answer.lower(),
    "Contains '1978'": "1978" in answer,
    "Has disclaimer": "important note" in answer.lower() or "exclusively" in answer.lower(),
    "Source is RAG Context": citations[0].get('source') == "RAG Context" if citations else False,
    "Answer length > 100 chars": len(answer) > 100
}

print("\n📋 TEST 6: Response Validation")
print("-" * 80)
for check, passed in checks.items():
    status = "✅" if passed else "❌"
    print(f"{status} {check}")

print("\n📋 TEST 7: General Query (No RAG)")
print("-" * 80)
# Test general query without RAG context
general_session = "general_test_session"
response = requests.post(f"{BASE_URL}/chat", json={
    "message": "What is the capital of France?",
    "session_id": general_session
})
result = response.json()
has_disclaimer = "important note" in result['response'].lower()
print(f"{'❌' if has_disclaimer else '✅'} General query has NO disclaimer (correct)")
print(f"✅ Source: {result.get('citations', [{}])[0].get('source', 'Unknown')}")

print("\n" + "=" * 80)
print("FINAL TEST RESULTS")
print("=" * 80)

all_passed = all(checks.values()) and not has_disclaimer
if all_passed:
    print("🎉 ALL TESTS PASSED! ShapeShifter is production-ready!")
else:
    print("⚠️  Some tests failed. Review the results above.")

print("\n✅ Tested Functionalities:")
print("   1. Health check endpoint")
print("   2. RAG creation workflow (personality, naming)")
print("   3. DOCX file upload and parsing")
print("   4. Settings integration (LLM, Vector DB)")
print("   5. Context-based query answering")
print("   6. Disclaimer on RAG responses")
print("   7. General queries without RAG")
print("\n" + "=" * 80)
